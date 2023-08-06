from typing import IO, List, Union

from google.cloud import translate
from google.api_core.exceptions import InvalidArgument

import docx

from getranet_automl.exceptions.exceptions import InvalidArgumentException

# Solve empty string Problem


class GoogleTranslationAPI:
    def __init__(
        self,
        project_id: str
    ) -> None:
        """
        Class for interacting with Google Translation API
        """
        self.client = translate.TranslationServiceAsyncClient()
        self.location: str = 'global'
        self.project_id: str = project_id
        self.parent: str = (f"projects/{self.project_id}/locations/"
                            + self.location)
        self.input: List[str] = []
        self.translation: List[str] = []
        self.num_chars: int = 0
        self.docx = None

    async def translate_text(
        self,
        input: Union[List[str], str],
        source_lang: str,
        target_lang: str
    ) -> List[str]:
        """Translating Text."""
        # Detail on supported types can be found here:
        # https://cloud.google.com/translate/docs/supported-formats
        self.input = [input] if isinstance(input, str) else input

        # skip empty strings
        input_list = list(filter(lambda a: a != '', self.input))
        try:
            response = await self.client.translate_text(
                request={
                    "parent": self.parent,
                    "contents": input_list,
                    "mime_type": "text/plain",  # text/plain, text/html
                    "source_language_code": source_lang,
                    "target_language_code": target_lang,
                }
            )
        except InvalidArgument as e:
            raise InvalidArgumentException(400, e)

        # Display the translation for each input text provided
        for translation in response.translations:
            self.translation.append(translation.translated_text)
        return self.translation

    async def convert_docx_to_list(
        self,
        file: IO
    ) -> List:
        self.docx = docx.Document(file)

        self.input = [paragraph.text for paragraph in self.docx.paragraphs]
        # list(filter(lambda a: a != '', input_list))
        self.num_chars = sum([len(text) for text in self.input])

        return self.input

    async def __convert_doc_to_docx(
        self
    ) -> docx.Document:
        ...

    async def translate_docx(
        self,
        source_lang: str,
        target_lang: str,
        output_file: str
    ) -> str:
        """Translate docx files."""
        # 1. batch translate paragraphs
        translation_list = await self.translate_text(
            self.input, source_lang, target_lang
        )
        # 2. replace paragraphs with translated paragraphs
        for i in range(len(self.docx.paragraphs)):
            original = self.docx.paragraphs[i].text
            # skip empty input strings
            if original != '':
                # take first translation
                translation = translation_list.pop(0)
                # add whitespaces at beginning or end of string if there where
                # any in the original
                translation = (translation + " " if original[-1] == " "
                               else translation)
                translation = (" " + translation if original[0] == " "
                               else translation)
                # replace original paragraph in docx with thranslation
                self.docx.paragraphs[i].text = translation

        self.docx.save(output_file)
        return self.docx
